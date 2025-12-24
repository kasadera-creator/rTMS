# rtms_app/services/sae_report.py
"""SAE Report generation for Brainsway using python-docx template."""
from __future__ import annotations
import os
import io
from typing import Optional, Dict, Any
from django.conf import settings
from docx import Document


def build_sae_context(session, sae_record=None) -> Dict[str, Any]:
    """Build context dict for SAE report from session and snapshot.

    Returns:
        dict with all placeholders for Word template
    """
    from rtms_app.models import SeriousAdverseEvent
    patient = session.patient
    snapshot = {}
    event_types = []
    other_text = ""

    if sae_record is None:
        sae_record = SeriousAdverseEvent.objects.filter(
            patient=patient,
            course_number=session.course_number,
            session=session,
        ).first()

    if sae_record:
        snapshot = sae_record.auto_snapshot or {}
        event_types = sae_record.event_types or []
        other_text = sae_record.other_text or ""

    # Event type mapping
    event_map = {
        "seizure": "けいれん発作",
        "finger_muscle": "手指の筋収縮",
        "syncope": "失神",
        "mania": "躁病・軽躁病の出現",
        "suicide_attempt": "自殺企図",
        "other": "その他",
    }
    event_labels = [event_map.get(e, e) for e in event_types]

    # Auto-fill fields
    context = {
        "FACILITY_NAME": "せいちりょう病院",  # update with your facility
        "FACILITY_PHONE": "（施設電話番号）",
        "REPORTER_NAME": "（担当者名）",
        "PHYSICIAN_COMMENT": "（医師コメント）",
        "PATIENT_INITIAL": _initial_from_name(patient.name) if patient.name else "（患者イニシャル）",
        "PATIENT_AGE": str(snapshot.get("age", patient.age)),
        "PATIENT_GENDER": snapshot.get("gender", patient.get_gender_display()),
        "DIAGNOSIS": snapshot.get("diagnosis", patient.diagnosis),
        "EVENT_DATE": snapshot.get("date", session.session_date.isoformat()),
        "EVENT_TYPES": "、".join(event_labels) if event_labels else "（有害事象名）",
        "EVENT_OTHER": other_text if other_text else "（その他詳細）",
        "MEDICATION": snapshot.get("medication_history", patient.medication_history) or "（併用薬）",
        "MT_VALUE": str(snapshot.get("mt_percent", "（MT値）")),
        "STIMULUS_INTENSITY": f"{snapshot.get('mt_percent', '（強度）')}%MT",
        "STIMULUS_SITE": snapshot.get("target_site", "左DLPFC"),
        "SESSION_COUNT": "（治療回数）",  # needs calculation or pass from caller
        "COIL_TYPE": snapshot.get("coil_type", "BrainsWay H1"),
        "FREQUENCY_HZ": str(snapshot.get("frequency_hz", "18.0")),
        "TRAIN_SECONDS": str(snapshot.get("train_seconds", "2.0")),
        "TRAIN_COUNT": str(snapshot.get("train_count", "55")),
        "TOTAL_PULSES": str(snapshot.get("total_pulses", "1980")),
        "INGESTION_STATUS": "（摂取状況）",
        "OUTCOME": "（転帰）",
        "OUTCOME_DATE": "（転帰日）",
        "EMAIL": "brainsway_saeinfo@cmi.co.jp",
    }

    return context


def _initial_from_name(name: str) -> str:
    """Extract patient initial (first char of each kanji word)."""
    parts = name.strip().split()
    if not parts:
        return "（患者イニシャル）"
    return "".join(p[0] for p in parts if p)


def render_sae_docx(template_path: str, context: Dict[str, Any]) -> bytes:
    """Render SAE report from template with context substitution.

    Args:
        template_path: path to .docx template with placeholders
        context: dict of {placeholder: value}

    Returns:
        bytes of generated Word document
    """
    doc = Document(template_path)

    # Replace in paragraphs
    for para in doc.paragraphs:
        for key, val in context.items():
            placeholder = f"{{{key}}}"
            if placeholder in para.text:
                # Simple inline replacement
                para.text = para.text.replace(placeholder, str(val))

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in context.items():
                    placeholder = f"{{{key}}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(val))

    # Write to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def get_missing_fields(context: Dict[str, Any]) -> list[str]:
    """Return list of placeholder keys that are still not filled (contain '（' or empty)."""
    missing = []
    for k, v in context.items():
        if not v or "（" in str(v):
            missing.append(k)
    return missing
